"""
Renders the AMRG markdown summary as a .docx (with embedded figures) and
uploads it to Google Drive as a Google Doc, via a service account.

Auth: set one of
  GOOGLE_SERVICE_ACCOUNT_FILE  - path to a service-account JSON key
  GOOGLE_SERVICE_ACCOUNT_JSON  - the JSON key contents inline (e.g. a CI secret)
and GOOGLE_DRIVE_FOLDER_ID (config.py) to an existing folder shared with
that service account's email (service accounts have no Drive storage quota
of their own -- the folder must belong to a real user/Shared Drive that has
granted the service account Editor access).
"""

import json
import os
import re

import config

DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.file"]
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GDOC_MIME = "application/vnd.google-apps.document"


class DriveUploadError(Exception):
    pass


def _markdown_to_docx(markdown_text: str, figures: list, title: str, out_path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required to build the summary document. Install it with "
            "`pip install python-docx`."
        ) from e

    doc = Document()
    doc.add_heading(title, level=0)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip() == "---":
            continue
        elif line.strip().startswith("- "):
            doc.add_paragraph(_strip_md_bold(line.strip()[2:]), style="List Bullet")
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.bold = True
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)

    if figures:
        doc.add_heading("Extracted Figures", level=1)
        for i, fig in enumerate(figures):
            path = fig.get("path")
            if not path or not os.path.exists(path):
                continue
            try:
                doc.add_picture(path, width=Inches(5.5))
            except Exception:
                continue
            caption = fig.get("caption") or f"Figure {i + 1} (page {fig.get('page', '?')})"
            cap_p = doc.add_paragraph()
            cap_run = cap_p.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9)

    doc.save(out_path)
    return out_path


def _strip_md_bold(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Split a line on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part.strip("*"))
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _get_drive_service():
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
    except ImportError as e:
        raise RuntimeError(
            "google-api-python-client and google-auth are required for Drive upload. "
            "Install them with `pip install google-api-python-client google-auth`."
        ) from e

    info = None
    if os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON"):
        info = json.loads(os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"])
    elif os.environ.get("GOOGLE_SERVICE_ACCOUNT_FILE"):
        with open(os.environ["GOOGLE_SERVICE_ACCOUNT_FILE"], "r", encoding="utf-8") as fh:
            info = json.load(fh)
    else:
        raise DriveUploadError(
            "No Google service-account credentials configured "
            "(set GOOGLE_SERVICE_ACCOUNT_JSON or GOOGLE_SERVICE_ACCOUNT_FILE)."
        )

    creds = service_account.Credentials.from_service_account_info(info, scopes=DRIVE_SCOPES)
    return build("drive", "v3", credentials=creds)


def upload_summary(markdown_text: str, figures: list, title: str, work_dir: str) -> dict:
    """Build the .docx and upload it to Drive as a Google Doc.

    Returns {"docx_path": str, "drive_file_id": str|None, "drive_url": str|None,
             "uploaded": bool, "error": str|None}.
    """
    os.makedirs(work_dir, exist_ok=True)
    safe_title = re.sub(r"[^\w\-. ]", "_", title)[:120]
    docx_path = os.path.join(work_dir, f"{safe_title}.docx")
    _markdown_to_docx(markdown_text, figures, title, docx_path)

    if not config.GOOGLE_DRIVE_FOLDER_ID:
        return {
            "docx_path": docx_path, "drive_file_id": None, "drive_url": None,
            "uploaded": False, "error": "GOOGLE_DRIVE_FOLDER_ID not configured; left as local .docx only.",
        }

    try:
        from googleapiclient.http import MediaFileUpload
        service = _get_drive_service()
        metadata = {
            "name": f"{title}.gdoc",
            "parents": [config.GOOGLE_DRIVE_FOLDER_ID],
            "mimeType": GDOC_MIME,
        }
        media = MediaFileUpload(docx_path, mimetype=DOCX_MIME, resumable=True)
        created = service.files().create(
            body=metadata, media_body=media, fields="id, webViewLink"
        ).execute()
        return {
            "docx_path": docx_path,
            "drive_file_id": created.get("id"),
            "drive_url": created.get("webViewLink"),
            "uploaded": True,
            "error": None,
        }
    except Exception as e:
        return {
            "docx_path": docx_path, "drive_file_id": None, "drive_url": None,
            "uploaded": False, "error": str(e),
        }
